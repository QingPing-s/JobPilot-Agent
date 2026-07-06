from io import BytesIO

import pymupdf
import pytest
from docx import Document

from src import document_parser
from src.ocr_service import OCRExtraction


def test_extract_markdown_removes_formatting():
    content = b"# Candidate\n\n**Skills:** Python, RAG\n\n- Built an Agent project"

    result = document_parser.extract_document_text(content, "resume.md")

    assert result.extraction_method == "markdown"
    assert "Candidate" in result.text
    assert "Python, RAG" in result.text
    assert "**" not in result.text


def test_extract_docx_includes_paragraphs_and_tables():
    document = Document()
    document.add_heading("Candidate", level=1)
    document.add_paragraph("Artificial Intelligence graduate student")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skills"
    table.rows[0].cells[1].text = "Python, RAG"
    buffer = BytesIO()
    document.save(buffer)

    result = document_parser.extract_document_text(buffer.getvalue(), "resume.docx")

    assert result.extraction_method == "docx"
    assert "Artificial Intelligence graduate student" in result.text
    assert "Skills | Python, RAG" in result.text


def test_extract_native_pdf_text():
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "Candidate Python RAG FastAPI LangGraph project experience")
    content = document.tobytes()
    document.close()

    result = document_parser.extract_document_text(content, "resume.pdf")

    assert result.extraction_method == "pdf_native"
    assert result.page_count == 1
    assert "Python RAG FastAPI" in result.text


def test_scanned_pdf_uses_ocr(monkeypatch):
    document = pymupdf.open()
    document.new_page()
    content = document.tobytes()
    document.close()
    monkeypatch.setattr(
        document_parser,
        "extract_image_text",
        lambda image_bytes, filename: OCRExtraction(
            text="Candidate Python RAG",
            lines=["Candidate Python RAG"],
            scores=[0.95],
            average_confidence=0.95,
        ),
    )

    result = document_parser.extract_document_text(content, "scan.pdf")

    assert result.extraction_method == "pdf_ocr"
    assert result.text == "Candidate Python RAG"
    assert result.average_confidence == pytest.approx(0.95)
    assert result.warnings


def test_rejects_legacy_doc():
    with pytest.raises(document_parser.DocumentParserError, match="DOCX"):
        document_parser.extract_document_text(b"legacy", "resume.doc")

